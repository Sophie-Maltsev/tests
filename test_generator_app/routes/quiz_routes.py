from flask import Blueprint, render_template, request, redirect, url_for, flash, session, send_file
from ..services import quiz_service, question_service  # Импортируем наши сервисы
from docx import Document
from docx.shared import Pt
from io import BytesIO

quiz_routes_bp = Blueprint('quiz_bp', __name__, template_folder='../templates')

@quiz_routes_bp.route('/')
def index_page():
    return render_template('index.html')


@quiz_routes_bp.route('/test/start', methods=['GET'])
def start_new_test_page():
    total_questions = question_service.count_total_questions()
    return render_template('start_test.html', total_questions=total_questions)


@quiz_routes_bp.route('/test/generate', methods=['POST'])
def generate_test_action():
    try:
        num_questions = int(request.form.get('num_questions', 5))
    except ValueError:
        flash('Некорректное количество вопросов.', 'error')
        return redirect(url_for('quiz_bp.start_new_test_page'))

    if num_questions <= 0:
        flash('Количество вопросов должно быть положительным.', 'error')
        return redirect(url_for('quiz_bp.start_new_test_page'))

    test_instance_id = quiz_service.generate_new_test_instance(num_questions)

    if test_instance_id:
        flash(f'Тест №{test_instance_id} успешно сгенерирован!', 'success')
        # Сохраняем ID теста в сессии, чтобы убедиться, что пользователь проходит именно этот тест
        session['current_test_instance_id'] = test_instance_id
        return redirect(url_for('quiz_bp.take_test_page', test_instance_id=test_instance_id))
    else:
        # Проверяем, есть ли вообще вопросы
        if question_service.count_total_questions() == 0:
            flash('В базе нет вопросов для генерации теста. Пожалуйста, добавьте вопросы.', 'error')
        else:
            flash(
                'Ошибка при генерации теста. Возможно, запрошено слишком много вопросов или произошла внутренняя ошибка.',
                'error')
        return redirect(url_for('quiz_bp.start_new_test_page'))


@quiz_routes_bp.route('/test/<int:test_instance_id>/take', methods=['GET'])
def take_test_page(test_instance_id):
    # Дополнительная проверка: пользователь должен проходить тот тест, который для него сгенерирован
    # Это простая проверка, можно усложнить при необходимости (например, если тесты назначаются пользователям)
    # stored_test_id = session.get('current_test_instance_id')
    # if stored_test_id != test_instance_id:
    #     flash('Ошибка: Попытка доступа к некорректному тесту.', 'error')
    #     return redirect(url_for('quiz_bp.start_new_test_page'))

    questions_for_test = quiz_service.get_test_questions_for_instance(test_instance_id)

    if questions_for_test is None:  # Сервис вернул None, значит тест не найден
        flash(f'Тест с ID {test_instance_id} не найден.', 'error')
        session.pop('current_test_instance_id', None)  # Очищаем из сессии
        return redirect(url_for('quiz_bp.start_new_test_page'))

    if not questions_for_test:  # Тест найден, но вопросов нет (маловероятно)
        flash('Для этого теста не найдено вопросов. Возможно, они были удалены.', 'error')
        session.pop('current_test_instance_id', None)
        return redirect(url_for('quiz_bp.start_new_test_page'))

    return render_template('take_test.html',
                           questions_with_answers=questions_for_test,
                           test_instance_id=test_instance_id)


@quiz_routes_bp.route('/test/<int:test_instance_id>/submit', methods=['POST'])
def submit_test_action(test_instance_id):
    # Проверка, что пользователь отправляет тот тест, который начал
    # stored_test_id = session.get('current_test_instance_id')
    # if stored_test_id != test_instance_id:
    #     flash('Ошибка: Попытка отправки результатов некорректного теста.', 'error')
    #     return redirect(url_for('quiz_bp.start_new_test_page'))

    user_submitted_answers = {}  # {question_id: selected_answer_id}
    for key, value in request.form.items():
        if key.startswith('question_'):
            try:
                # Преобразуем ID вопроса и ответа в int
                question_id = int(key.split('_')[1])
                selected_answer_id = int(value)
                user_submitted_answers[question_id] = selected_answer_id
            except (ValueError, IndexError):
                flash('Обнаружены некорректные данные в отправленных ответах.', 'error')
                return redirect(url_for('quiz_bp.take_test_page', test_instance_id=test_instance_id))

    if not user_submitted_answers:
        flash('Вы не ответили ни на один вопрос.', 'warning')
        return redirect(url_for('quiz_bp.take_test_page', test_instance_id=test_instance_id))

    evaluation_result = quiz_service.submit_and_evaluate_test(test_instance_id, user_submitted_answers)

    if evaluation_result and 'error' not in evaluation_result:
        flash('Тест успешно завершен!', 'success')
        session.pop('current_test_instance_id', None)  # Очищаем ID теста из сессии после успешной сдачи
        return render_template('test_results.html', attempt=evaluation_result)
    else:
        error_message = evaluation_result.get('error',
                                              'Произошла неизвестная ошибка при проверке теста.') if evaluation_result else 'Не удалось обработать результаты теста.'
        flash(error_message, 'error')
        return redirect(url_for('quiz_bp.take_test_page', test_instance_id=test_instance_id))

@quiz_routes_bp.route('/test/<int:test_instance_id>/download', methods=['GET'])
def download_test_docx(test_instance_id):
    """
    Генерирует тест в формате .docx для распечатки
    """
    questions = quiz_service.get_test_questions_for_instance(test_instance_id)
    if not questions:
        flash("Невозможно найти вопросы для этого теста.", "error")
        return redirect(url_for('quiz_bp.start_new_test_page'))

    doc = Document()
    doc.add_heading(f'Тест №{test_instance_id}', 0)

    for idx, q in enumerate(questions, start=1):
        p = doc.add_paragraph(f"{idx}. {q['question_text']}")
        p.runs[0].font.size = Pt(12)
        if q.get('topic'):
            doc.add_paragraph(f"(Тема: {q['topic']})", style="Intense Quote")
        for a_idx, ans in enumerate(q['answers'], start=1):
            doc.add_paragraph(f"   {chr(96+a_idx)}) {ans['answer_text']}", style='List Bullet')

        doc.add_paragraph("")  # Пустая строка между вопросами

    # Сохраняем в память
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    filename = f"test_{test_instance_id}.docx"
    return send_file(
        doc_io,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )